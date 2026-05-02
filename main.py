from fastapi import FastAPI, File, UploadFile, Form, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
import uvicorn
import pandas as pd
import io
import docx
import tempfile
import os
from translation_service import TMTService

import fitz  # PyMuPDF

# Initialize services
app = FastAPI(
    title="FastAPI Backend",
    description="A production-ready FastAPI structure for document processing and translation.",
    version="1.1.0"
)
tmt_service = TMTService()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/ping", tags=["Health Check"])
async def ping():
    return {"status": "ok", "message": "pong"}

@app.post("/translate/csv", tags=["Translation"])
async def translate_csv(
    file: UploadFile = File(...),
    target_lang: str = Form(...),
    source_lang: str = Form("en")
):
    """
    Accepts a CSV file, translates all string content, and returns the translated CSV.
    """
    # Read the CSV into a pandas DataFrame
    contents = await file.read()
    df = pd.read_csv(io.BytesIO(contents))

    # Define a helper to translate if the value is a string
    def translate_value(val):
        if isinstance(val, str):
            return tmt_service.translate_text(val, source_lang, target_lang)
        return val

    # Apply translation across the entire dataframe
    translated_df = df.map(translate_value)

    # Convert translated dataframe back to CSV in memory
    stream = io.StringIO()
    translated_df.to_csv(stream, index=False)
    
    # Reset stream position and return as a downloadable file
    response = StreamingResponse(
        iter([stream.getvalue()]),
        media_type="text/csv"
    )
    response.headers["Content-Disposition"] = f"attachment; filename=translated_{file.filename}"
    
    return response

@app.post("/translate/docx", tags=["Translation"])
async def translate_docx(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    target_lang: str = Form(...),
    source_lang: str = Form("en")
):
    """
    Accepts a .docx file, translates text while preserving structure/formatting, and returns the modified file.
    """
    # Read docx from upload
    contents = await file.read()
    doc = docx.Document(io.BytesIO(contents))

    # Helper to translate runs in a list of paragraphs
    def translate_paragraphs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                if run.text.strip():
                    run.text = tmt_service.translate_text(run.text, source_lang, target_lang)

    # Translate main document paragraphs
    translate_paragraphs(doc.paragraphs)

    # Translate tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                translate_paragraphs(cell.paragraphs)

    # Save to a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    temp_file.close()

    # Clean up the temp file after response is sent
    background_tasks.add_task(os.remove, temp_file.name)

    return FileResponse(
        temp_file.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"translated_{file.filename}"
    )

# ... keep your other imports

@app.post("/translate/pdf", tags=["Translation"])
async def translate_pdf(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    target_lang: str = Form(...),
    source_lang: str = Form(...)
):
    contents = await file.read()
    try:
        # Open the original PDF and create a new blank PDF
        doc = fitz.open(stream=contents, filetype="pdf")
        out_doc = fitz.open()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid PDF file: {e}")
    
    current_dir = os.path.dirname(os.path.abspath(__file__))
    font_path = os.path.join(current_dir, "nepalifront.ttf") 

    for page in doc:
        # Create a new page with exactly the same size
        new_page = out_doc.new_page(width=page.rect.width, height=page.rect.height)
        
        # Check if the font file exists for Nepali/Tamang
        font_registered = False
        if os.path.exists(font_path) and target_lang in ["ne", "taj"]:
            new_page.insert_font(fontfile=font_path, fontname="devanagari")
            font_registered = True

        blocks = page.get_text("blocks")
        for b in blocks:
            # b[4] is the text content
            original_text = b[4].replace('\n', ' ').strip()
            # b[:4] gives the (x0, y0, x1, y1) coordinates
            rect = fitz.Rect(b[:4])
            
            if original_text:
                # Call your translation service
                translated_text = tmt_service.translate_text(original_text, source_lang, target_lang)
                
                try:
                    if target_lang in ["ne", "taj"] and font_registered:
                        # Use HTMLBox for better Devanagari rendering
                        html = f"<div style='font-family:devanagari; font-size:10pt; color:black;'>{translated_text}</div>"
                        new_page.insert_htmlbox(rect, html)
                    else:
                        # Standard English text insertion
                        new_page.insert_textbox(
                            rect, 
                            translated_text, 
                            fontsize=10, 
                            fontname="helv", 
                            color=(0, 0, 0) # Explicitly set color to BLACK
                        )
                except Exception:
                    # Final fallback: just put the text at the starting point if the box fails
                    new_page.insert_text(
                        rect.tl, 
                        translated_text, 
                        fontsize=10, 
                        fontname="devanagari" if font_registered else "helv",
                        color=(0, 0, 0)
                    )

    # Save to a temporary file
    fd, temp_path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd) 

    try:
        out_doc.save(temp_path, garbage=3, deflate=True)
        out_doc.close()
        doc.close()
    except Exception as e:
        if os.path.exists(temp_path): os.remove(temp_path)
        raise HTTPException(status_code=500, detail=f"Save Error: {e}")

    background_tasks.add_task(os.remove, temp_path)
    return FileResponse(temp_path, filename=f"translated_{file.filename}")

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)